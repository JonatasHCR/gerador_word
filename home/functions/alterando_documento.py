import docx#pacote para abir o documento word, salvar, editar
from pathlib import Path

CAMINHO_MODELO = Path(__file__).parent / 'modelos'
'''
Local para ja com os dados tratados, usar o modelo para substituir, lembrando 
que tera que criar uma pasta chamada 'modelos', ela foi barrada no gitignore
mas é nessa pasta functions
'''

def substituindo_texto(paragraph, subst_text_list: list):
    #essa função recebe a lista que dentro tem uma tupla com chave e valor
    #essa chave representa as palavras que estão dentro do arquivo e serão 
    #substituídas pelos valores
    for old_text in subst_text_list:
        new_text = old_text[1]
        if old_text[0] in paragraph.text:
            #feito isso temos que preserva a formatação que a palavra se encontra
            linha = paragraph.runs
            for item in linha:
                #nao vamos alterar tudo por isso pegamos apenas o texto e 
                # preservarmos a formatação
                if old_text[0] in item.text:
                    item.text = item.text.replace(old_text[0], new_text)

def gerando_documento(variables_values: list, caminho_modelo):
    for tuplas in variables_values:
        match tuplas[0]:
            case 'NOME_MODELO':
                path_file_model = tuplas[1].strip()
            case 'CAMINHO_SAIDA':
                caminho_saida = tuplas[1]
            case 'ARQUIVO_SAIDA':
                path_file_save = tuplas[1].strip() + '.doc'
        continue
    caminho_modelo = caminho_modelo / path_file_model
    caminho_saida = f'{caminho_saida}\\{path_file_save}'
    try:#ainda em processo de ajeitar para melhorar o except
        doc = docx.Document(caminho_modelo)
        for linha in doc.paragraphs:
                substituindo_texto(linha,variables_values)
        #depois de tudo ele usa o caminho de saída fornecido pelo usuário
        doc.save(caminho_saida)
    except:
        print('Não foi possível salvar o arquivo')#caso de erro(lembrando ta 
        #em processo de criação vou melhorar os erros)

def gerando_modelo(nome_modelo):
    nome_modelo = nome_modelo + '.doc'
    caminho_modelo = CAMINHO_MODELO / nome_modelo
    doc = docx.Document()
    doc.save(caminho_modelo)